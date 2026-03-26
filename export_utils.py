from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def export_story_to_docx(story, tags, characters, world_notes, chapters, output_dir="exports"):
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    safe_name = story["title"].replace(" ", "_").replace("/", "_")
    file_path = output_path / f"{safe_name}.docx"

    doc = Document()
    doc.add_heading(story["title"], level=0)

    doc.add_paragraph(f"Genre: {story['genre'] or 'N/A'}")
    doc.add_paragraph(f"Summary: {story['summary'] or 'N/A'}")

    doc.add_heading("Tags", level=1)
    if tags:
        doc.add_paragraph(", ".join(tag["tag_name"] for tag in tags))
    else:
        doc.add_paragraph("No tags available.")

    doc.add_heading("Characters", level=1)
    if characters:
        for c in characters:
            doc.add_heading(c["name"], level=2)
            doc.add_paragraph(f"Role: {c['role'] or 'N/A'}")
            doc.add_paragraph(f"Traits: {c['traits'] or 'N/A'}")
            doc.add_paragraph(f"Goals: {c['goals'] or 'N/A'}")
            doc.add_paragraph(f"Description: {c['description'] or 'N/A'}")
    else:
        doc.add_paragraph("No character profiles available.")

    doc.add_heading("World Notes", level=1)
    if world_notes:
        for note in world_notes:
            doc.add_heading(f"{note['title']} [{note['category'] or 'General'}]", level=2)
            doc.add_paragraph(note["content"])
    else:
        doc.add_paragraph("No world notes available.")

    doc.add_heading("Chapters", level=1)
    if chapters:
        for ch in chapters:
            doc.add_heading(ch["chapter_title"] or "Untitled Chapter", level=2)
            doc.add_paragraph(ch["content"])
    else:
        doc.add_paragraph("No chapters available.")

    doc.save(file_path)
    return str(file_path)


def export_story_to_pdf(story, tags, characters, world_notes, chapters, output_dir="exports"):
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    safe_name = story["title"].replace(" ", "_").replace("/", "_")
    file_path = output_path / f"{safe_name}.pdf"

    c = canvas.Canvas(str(file_path), pagesize=letter)
    width, height = letter
    y = height - 50

    def write_line(text, line_gap=15):
        nonlocal y
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, text[:110])
        y -= line_gap

    write_line(story["title"], 20)
    write_line(f"Genre: {story['genre'] or 'N/A'}")
    write_line(f"Summary: {story['summary'] or 'N/A'}")
    y -= 10

    write_line("Tags:", 20)
    write_line(", ".join(tag["tag_name"] for tag in tags) if tags else "No tags available.")
    y -= 10

    write_line("Characters:", 20)
    if characters:
        for char in characters:
            write_line(f"Name: {char['name']}")
            write_line(f"Role: {char['role'] or 'N/A'}")
            write_line(f"Traits: {char['traits'] or 'N/A'}")
            write_line(f"Goals: {char['goals'] or 'N/A'}")
            write_line(f"Description: {char['description'] or 'N/A'}")
            y -= 5
    else:
        write_line("No character profiles available.")
    y -= 10

    write_line("World Notes:", 20)
    if world_notes:
        for note in world_notes:
            write_line(f"{note['title']} [{note['category'] or 'General'}]")
            for line in (note["content"] or "").splitlines():
                write_line(line)
            y -= 5
    else:
        write_line("No world notes available.")
    y -= 10

    write_line("Chapters:", 20)
    if chapters:
        for ch in chapters:
            write_line(ch["chapter_title"] or "Untitled Chapter")
            for line in (ch["content"] or "").splitlines():
                write_line(line)
            y -= 8
    else:
        write_line("No chapters available.")

    c.save()
    return str(file_path)